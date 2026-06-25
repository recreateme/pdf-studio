"""
PDF Studio - OCR 识别引擎
支持 RapidOCR（主选）与 PaddleOCR（备选）
完全本地离线运行
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Optional, Protocol

import fitz
from PIL import Image

from app.utils.logger import logger
from app.utils.helpers import make_temp_file, make_temp_dir, get_unique_path


# ─────────────────────────────────────────────
# 数据类型
# ─────────────────────────────────────────────

@dataclass
class OCRBlock:
    """OCR识别块（单行文字）"""
    text: str
    confidence: float
    bbox: tuple[float, float, float, float]   # (x0,y0,x1,y1) 归一化坐标
    page_index: int = 0


@dataclass
class OCRResult:
    """单页OCR结果"""
    page_index: int
    blocks: list[OCRBlock] = field(default_factory=list)
    full_text: str = ""
    processing_time: float = 0.0

    def get_text(self) -> str:
        return "\n".join(b.text for b in self.blocks)


@dataclass
class OCROptions:
    """OCR选项"""
    languages: list[str] = field(default_factory=lambda: ["ch", "en"])
    use_gpu: bool = False
    confidence_threshold: float = 0.5
    dpi: int = 200                   # 渲染DPI（越高越准确）
    pages: Optional[list[int]] = None    # None=全部
    output_format: str = "txt"       # txt/docx/markdown/json/searchable_pdf
    output_path: Optional[Path] = None


# ─────────────────────────────────────────────
# OCR 引擎协议（接口）
# ─────────────────────────────────────────────

class OCREngineProtocol(Protocol):
    """OCR引擎统一接口协议"""

    def is_available(self) -> bool:
        """检查引擎是否可用"""
        ...

    def recognize_image(
        self,
        img: Image.Image,
        languages: list[str],
    ) -> list[OCRBlock]:
        """识别单张图片，返回OCR块列表"""
        ...


# ─────────────────────────────────────────────
# RapidOCR 引擎实现
# ─────────────────────────────────────────────

class RapidOCREngine:
    """
    RapidOCR 引擎（推荐，基于ONNX，轻量快速）

    安装：pip install rapidocr-onnxruntime
    """

    def __init__(self, ocr_settings=None) -> None:
        self._ocr_settings = ocr_settings
        self._engine = None

    def is_available(self) -> bool:
        """检查 rapidocr_onnxruntime 是否已安装"""
        try:
            import rapidocr_onnxruntime
            return True
        except ImportError:
            return False

    def _build_rapidocr_kwargs(self) -> dict:
        """根据 OCR 设置构建 RapidOCR 初始化参数"""
        if not self._ocr_settings:
            return {}

        s = self._ocr_settings
        kwargs: dict = {}

        path_map = (
            ("det_model_path", "det_model_path"),
            ("rec_model_path", "rec_model_path"),
            ("cls_model_path", "cls_model_path"),
        )
        for attr, key in path_map:
            model_path = getattr(s, attr, "").strip()
            if model_path and Path(model_path).is_file():
                kwargs[key] = model_path

        if s.use_gpu:
            kwargs["det_use_cuda"] = True
            kwargs["rec_use_cuda"] = True
            kwargs["cls_use_cuda"] = True

        return kwargs

    def _get_engine(self):
        if self._engine is None:
            from rapidocr_onnxruntime import RapidOCR

            kwargs = self._build_rapidocr_kwargs()
            self._engine = RapidOCR(**kwargs)
            logger.info("RapidOCR 引擎初始化完成")
        return self._engine

    def recognize_image(
        self,
        img: Image.Image,
        languages: list[str] = None,
    ) -> list[OCRBlock]:
        """
        识别图片中的文字

        Args:
            img: PIL Image
            languages: 语言列表（RapidOCR自动检测，此参数预留）

        Returns:
            OCRBlock列表
        """
        import numpy as np
        engine = self._get_engine()
        arr = np.array(img)
        result, _ = engine(arr)

        if result is None:
            return []

        blocks = []
        for item in result:
            # item: [[x0,y0],[x1,y1],[x2,y2],[x3,y3]], text, confidence
            box, text, conf = item
            xs = [p[0] for p in box]
            ys = [p[1] for p in box]
            w, h = img.size
            bbox = (
                min(xs) / w, min(ys) / h,
                max(xs) / w, max(ys) / h,
            )
            blocks.append(OCRBlock(
                text=text,
                confidence=float(conf),
                bbox=bbox,
            ))

        return blocks


# ─────────────────────────────────────────────
# PaddleOCR 引擎实现（备选）
# ─────────────────────────────────────────────

class PaddleOCREngine:
    """
    PaddleOCR 引擎（备选，精度更高，体积较大）

    安装：pip install paddlepaddle paddleocr
    """

    LANG_MAP = {
        "ch": "ch", "en": "en",
        "japan": "japan", "korean": "korean",
        "french": "fr", "german": "german",
    }

    def __init__(self, ocr_settings=None) -> None:
        self._ocr_settings = ocr_settings
        self._engines: dict[str, object] = {}

    def is_available(self) -> bool:
        try:
            import paddleocr
            return True
        except ImportError:
            return False

    def _get_engine(self, lang: str = "ch"):
        if lang not in self._engines:
            from paddleocr import PaddleOCR

            use_gpu = bool(self._ocr_settings and self._ocr_settings.use_gpu)
            self._engines[lang] = PaddleOCR(
                use_angle_cls=True,
                lang=self.LANG_MAP.get(lang, "ch"),
                use_gpu=use_gpu,
                show_log=False,
            )
            logger.info(f"PaddleOCR 引擎初始化完成 (lang={lang}, gpu={use_gpu})")
        return self._engines[lang]

    def recognize_image(
        self,
        img: Image.Image,
        languages: list[str] = None,
    ) -> list[OCRBlock]:
        import numpy as np
        lang = (languages or ["ch"])[0]
        engine = self._get_engine(lang)
        arr = np.array(img)
        result = engine.ocr(arr, cls=True)

        if not result or not result[0]:
            return []

        blocks = []
        w, h = img.size
        for line in result[0]:
            box = line[0]   # [[x0,y0],[x1,y1],[x2,y2],[x3,y3]]
            text = line[1][0]
            conf = float(line[1][1])
            xs = [p[0] for p in box]
            ys = [p[1] for p in box]
            bbox = (
                min(xs) / w, min(ys) / h,
                max(xs) / w, max(ys) / h,
            )
            blocks.append(OCRBlock(text=text, confidence=conf, bbox=bbox))

        return blocks


# ─────────────────────────────────────────────
# OCR 管理器（统一入口）
# ─────────────────────────────────────────────

class OCRManager:
    """
    OCR 管理器
    自动选择可用引擎，提供统一的PDF/图片OCR接口
    """

    def __init__(self, engine_name: str = "auto") -> None:
        self._engine = self._init_engine(engine_name)

    def _init_engine(self, name: str) -> OCREngineProtocol:
        """初始化OCR引擎"""
        from app.config.settings import settings_mgr

        ocr_cfg = settings_mgr.ocr

        if name == "rapidocr" or name == "auto":
            engine = RapidOCREngine(ocr_cfg)
            if engine.is_available():
                logger.info("使用 RapidOCR 引擎")
                return engine

        if name == "paddleocr" or name == "auto":
            engine = PaddleOCREngine(ocr_cfg)
            if engine.is_available():
                logger.info("使用 PaddleOCR 引擎")
                return engine

        logger.warning("未找到可用的OCR引擎，OCR功能将不可用")
        return None

    @property
    def is_available(self) -> bool:
        return self._engine is not None

    # ── PDF OCR ───────────────────────────────

    def ocr_pdf(
        self,
        path: str | Path,
        options: OCROptions,
        progress_cb: Optional[Callable[[int, int], None]] = None,
        should_cancel: Optional[Callable[[], bool]] = None,
    ) -> list[OCRResult]:
        """
        对PDF执行OCR

        Args:
            path: PDF文件路径
            options: OCR选项
            progress_cb: 进度回调

        Returns:
            每页OCR结果列表
        """
        if not self.is_available:
            raise RuntimeError("OCR引擎不可用，请安装 rapidocr-onnxruntime")

        path = Path(path)
        doc = fitz.open(str(path))
        total_pages = len(doc)
        target_pages = options.pages if options.pages is not None else list(range(total_pages))
        results = []

        import time
        for idx, page_idx in enumerate(target_pages):
            if should_cancel and should_cancel():
                break
            if page_idx >= total_pages:
                continue

            t0 = time.time()
            page = doc[page_idx]
            mat = fitz.Matrix(options.dpi / 72, options.dpi / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

            blocks = self._engine.recognize_image(img, options.languages)

            # 过滤低置信度结果
            blocks = [b for b in blocks if b.confidence >= options.confidence_threshold]
            for b in blocks:
                b.page_index = page_idx

            result = OCRResult(
                page_index=page_idx,
                blocks=blocks,
                full_text="\n".join(b.text for b in blocks),
                processing_time=time.time() - t0,
            )
            results.append(result)
            logger.debug(f"OCR 第 {page_idx+1} 页: {len(blocks)} 个文字块, {result.processing_time:.2f}s")

            if progress_cb:
                progress_cb(idx + 1, len(target_pages))

        doc.close()
        return results

    # ── 图片 OCR ─────────────────────────────

    def ocr_image(
        self,
        path: str | Path,
        options: OCROptions,
    ) -> OCRResult:
        """对单张图片执行OCR"""
        if not self.is_available:
            raise RuntimeError("OCR引擎不可用")

        import time
        t0 = time.time()
        img = Image.open(str(path)).convert("RGB")
        blocks = self._engine.recognize_image(img, options.languages)
        blocks = [b for b in blocks if b.confidence >= options.confidence_threshold]

        return OCRResult(
            page_index=0,
            blocks=blocks,
            full_text="\n".join(b.text for b in blocks),
            processing_time=time.time() - t0,
        )

    # ── 结果导出 ──────────────────────────────

    def export_results(
        self,
        results: list[OCRResult],
        output_path: Path,
        format: str = "txt",
    ) -> Path:
        """
        导出OCR结果

        Args:
            results: OCR结果列表
            output_path: 输出路径（不含后缀）
            format: txt/docx/markdown/json/searchable_pdf

        Returns:
            实际输出文件路径
        """
        format = format.lower()
        exporters = {
            "txt": self._export_txt,
            "docx": self._export_docx,
            "markdown": self._export_markdown,
            "json": self._export_json,
        }

        if format in exporters:
            return exporters[format](results, output_path)
        else:
            raise ValueError(f"不支持的导出格式: {format}")

    def _export_txt(self, results: list[OCRResult], base_path: Path) -> Path:
        out = base_path.with_suffix(".txt")
        lines = []
        for r in sorted(results, key=lambda x: x.page_index):
            lines.append(f"=== 第 {r.page_index + 1} 页 ===")
            lines.append(r.full_text)
            lines.append("")
        out.write_text("\n".join(lines), encoding="utf-8")
        logger.info(f"OCR结果已导出: {out.name}")
        return out

    def _export_markdown(self, results: list[OCRResult], base_path: Path) -> Path:
        out = base_path.with_suffix(".md")
        lines = []
        for r in sorted(results, key=lambda x: x.page_index):
            lines.append(f"## 第 {r.page_index + 1} 页\n")
            lines.append(r.full_text)
            lines.append("\n---\n")
        out.write_text("\n".join(lines), encoding="utf-8")
        return out

    def _export_json(self, results: list[OCRResult], base_path: Path) -> Path:
        out = base_path.with_suffix(".json")
        data = [
            {
                "page": r.page_index + 1,
                "full_text": r.full_text,
                "blocks": [
                    {"text": b.text, "confidence": b.confidence, "bbox": b.bbox}
                    for b in r.blocks
                ],
            }
            for r in sorted(results, key=lambda x: x.page_index)
        ]
        out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return out

    def _export_docx(self, results: list[OCRResult], base_path: Path) -> Path:
        """导出为 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            logger.warning("python-docx未安装，降级为TXT输出")
            return self._export_txt(results, base_path)

        out = base_path.with_suffix(".docx")
        doc = Document()
        doc.add_heading("OCR 识别结果", 0)

        for r in sorted(results, key=lambda x: x.page_index):
            doc.add_heading(f"第 {r.page_index + 1} 页", level=1)
            for block in r.blocks:
                p = doc.add_paragraph(block.text)
                p.runs[0].font.size = Pt(11)
            doc.add_paragraph()

        doc.save(str(out))
        logger.info(f"OCR结果已导出Word: {out.name}")
        return out

    def generate_searchable_pdf(
        self,
        original_pdf: str | Path,
        results: list[OCRResult],
        output_path: str | Path,
        should_cancel: Optional[Callable[[], bool]] = None,
        cleanup_on_cancel: bool = False,
    ) -> Path:
        """
        生成可搜索PDF（在原PDF上叠加透明文字层）

        [接口已定义，完整实现使用 fitz text insertion]
        """
        original_pdf = Path(original_pdf)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 为了提升中文文本层的可搜索性，尽量显式指定一个中文字体文件。
        font_file = None
        try:
            candidates = [
                Path(r"C:\Windows\Fonts\simhei.ttf"),
                Path(r"C:\Windows\Fonts\simsun.ttc"),
                Path(r"C:\Windows\Fonts\simsun.ttf"),
                Path(r"C:\Windows\Fonts\simfang.ttf"),
            ]
            for c in candidates:
                if c.exists():
                    font_file = str(c)
                    break
        except Exception:
            font_file = None

        doc = fitz.open(str(original_pdf))
        result_map = {r.page_index: r for r in results}

        cancelled = False
        for i, page in enumerate(doc):
            if should_cancel and should_cancel():
                cancelled = True
                break
            if i not in result_map:
                continue
            result = result_map[i]
            rect = page.rect

            for block in result.blocks:
                # 将归一化坐标转回页面坐标
                x0 = block.bbox[0] * rect.width
                y0 = block.bbox[1] * rect.height
                x1 = block.bbox[2] * rect.width
                y1 = block.bbox[3] * rect.height
                # 插入透明文字（opacity=0 不可见但可搜索）
                page.insert_text(
                    fitz.Point(x0, y1),
                    block.text,
                    fontsize=max(8, (y1 - y0) * 0.8),
                    # PyMuPDF 对 fill_opacity=0.0 时可能导致文本层不可提取
                    # 使用极小不透明度保证“几乎不可见”且仍可被搜索/抽取。
                    color=(0, 0, 0),
                    fill_opacity=0.01,
                    fontfile=font_file,
                )

        if not cancelled:
            doc.save(str(output_path))
        doc.close()

        if cancelled and cleanup_on_cancel:
            try:
                Path(output_path).unlink(missing_ok=True)
            except Exception:
                pass
        logger.info(f"可搜索PDF生成完成: {output_path.name}")
        return output_path
